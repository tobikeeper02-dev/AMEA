"""Utilities to export AMEA findings into client-ready formats."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from ..pipeline import ComparativeAnalysis, MarketAnalysisResult


def export_to_docx(analysis: ComparativeAnalysis, path: Path) -> Path:
    """Create a Microsoft Word document summarizing the findings."""
    document = Document()
    document.add_heading(f"{analysis.company} Market Entry Analysis", level=1)
    document.add_paragraph(f"Industry: {analysis.industry}")
    document.add_paragraph(f"Engagement focus: {analysis.use_case}")

    if best := analysis.best_market():
        document.add_paragraph(f"Recommended market: {best.country} (score {best.score.composite}/100)")

    if analysis.company_brief:
        document.add_heading("Company & Industry Intelligence", level=2)
        summary = analysis.company_brief.get("profile_summary")
        if summary:
            document.add_paragraph(summary)

        sections = [
            ("Strategic fit", analysis.company_brief.get("strategic_fit", [])),
            ("Demand drivers", analysis.company_brief.get("demand_drivers", [])),
            ("Technology enablers", analysis.company_brief.get("technology_enablers", [])),
            ("Regulatory watch", analysis.company_brief.get("regulatory_watch", [])),
            ("Sustainability factors", analysis.company_brief.get("sustainability_factors", [])),
            ("Risk watch", analysis.company_brief.get("risk_watch", [])),
        ]
        for heading, bullets in sections:
            cleaned = [bullet for bullet in bullets if bullet]
            if not cleaned:
                continue
            document.add_paragraph(heading, style="List Bullet")
            for bullet in cleaned:
                document.add_paragraph(bullet, style="List Number")

    for market in analysis.markets:
        document.add_heading(market.country, level=2)
        document.add_paragraph(f"Composite opportunity score: {market.score.composite}/100")
        document.add_paragraph(f"Preferred entry mode: {market.entry_mode}")

        document.add_heading("PESTEL Highlights", level=3)
        for dimension, bullets in market.pestel.items():
            document.add_paragraph(dimension, style="List Bullet")
            for bullet in bullets:
                document.add_paragraph(bullet, style="List Number")

        if market.news:
            document.add_heading("Recent Signals", level=3)
            for headline in market.news:
                document.add_paragraph(headline, style="List Bullet")

        if market.turnaround_actions:
            document.add_heading("Risk Mitigations", level=3)
            for theme, action in market.turnaround_actions.items():
                document.add_paragraph(f"{theme.title()}: {action}", style="List Bullet")

        if market.sources:
            document.add_heading("Sources", level=3)
            for source in market.sources:
                document.add_paragraph(source, style="List Bullet")

    document.add_page_break()
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path
